# utilities.py
import re
import os
from fastapi import Form
import re
import io
from io import BytesIO
import base64
import math
import paramiko
from dateutil import parser
from pathlib import Path
import filetype
from PyPDF2 import PdfReader
from better_profanity import profanity
from adblockparser import AdblockRules
from pdfminer.high_level import extract_text as pdfminer_extract_text
import pytesseract
from pdf2image import convert_from_path
from datetime import datetime, date, time
from typing import Optional, Union, Tuple, List
import asyncio
from PIL import Image
import pandas as pd
import ipaddress
from scipy.spatial import cKDTree
from dateutil.parser import parse as dateutil_parse
from docx import Document
import aiohttp
from bs4 import BeautifulSoup
from readability import Document as ReadabilityDocument
from markdownify import markdownify as md
from sshtunnel import SSHTunnelForwarder
from urllib.parse import urlparse
from fastapi import Depends, HTTPException, Request, UploadFile
from fastapi.security.api_key import APIKeyHeader

from sijapi import L, API, Archivist, YEAR_FMT, MONTH_FMT, DAY_FMT, DAY_SHORT_FMT, OBSIDIAN_VAULT_DIR, ALLOWED_FILENAME_CHARS, MAX_PATH_LENGTH, ARCHIVE_DIR

logger = L.get_module_logger('utilities')
def debug(text: str): logger.debug(text)
def info(text: str): logger.info(text)
def warn(text: str): logger.warning(text)
def err(text: str): logger.error(text)
def crit(text: str): logger.critical(text)

api_key_header = APIKeyHeader(name="Authorization", auto_error=False)

def validate_api_key(request: Request, api_key: str = Depends(api_key_header)):
    if request.url.path in API.PUBLIC:
        return

    client_ip = ipaddress.ip_address(request.client.host)
    trusted_subnets = [ipaddress.ip_network(subnet) for subnet in API.TRUSTED_SUBNETS]
    if any(client_ip in subnet for subnet in trusted_subnets):
        return

    # Check header-based API key
    if api_key:
        if api_key.lower().startswith("bearer "):
            api_key = api_key.lower().split("bearer ")[-1]
        if api_key in API.KEYS:
            return

    # Check query-based API key
    api_key_query = request.query_params.get("api_key")
    if api_key_query in API.KEYS:
        return

    raise HTTPException(status_code=401, detail="Invalid or missing API key")


def assemble_archive_path(filename: str, extension: str = None, date_time: datetime = datetime.now(), subdir: str = None) -> Tuple[Path, Path]:
    year = date_time.strftime(YEAR_FMT)
    month = date_time.strftime(MONTH_FMT)
    day = date_time.strftime(DAY_FMT)
    day_short = date_time.strftime(DAY_SHORT_FMT)
    timestamp = date_time.strftime("%H%M%S")
    
    # Handle extension priority
    base_name, original_ext = os.path.splitext(filename)
    
    if extension is not None:
        # Use the provided extension parameter
        final_extension = extension if extension.startswith('.') else f'.{extension}'
    elif original_ext:
        # Use the original file extension if present
        final_extension = original_ext
    else:
        # Default to ".md" if no extension is provided or present
        final_extension = ".md"
    
    # Initial sanitization
    sanitized_base = sanitize_filename(base_name, '')
    filename = f"{day_short} {timestamp} {sanitized_base}{final_extension}"
    
    relative_path = Path(year) / month / day / filename
    absolute_path = Archivist.dir / relative_path
    
    # Ensure the total path length doesn't exceed MAX_PATH_LENGTH
    while len(str(absolute_path)) > MAX_PATH_LENGTH and len(sanitized_base) > 0:
        sanitized_base = sanitized_base[:-1]
        filename = f"{day_short} {timestamp} {sanitized_base}{final_extension}"
        relative_path = Path(year) / month / day / filename
        absolute_path = ARCHIVE_DIR / relative_path
    
    # If we've exhausted sanitized_base and the path is still too long
    if len(str(absolute_path)) > MAX_PATH_LENGTH:
        # Use a hash of the original filename to ensure uniqueness
        hash_suffix = hashlib.md5(base_name.encode()).hexdigest()[:8]
        filename = f"{day_short} {timestamp} {hash_suffix}{final_extension}"
        relative_path = Path(year) / month / day / filename
        absolute_path = ARCHIVE_DIR / relative_path
    
    # Final check and truncation if necessary
    if len(str(absolute_path)) > MAX_PATH_LENGTH:
        overflow = len(str(absolute_path)) - MAX_PATH_LENGTH
        absolute_path = Path(str(absolute_path)[:-overflow])
        relative_path = Path(str(relative_path)[:-overflow])
    
    return absolute_path, relative_path



def assemble_journal_path(date_time: datetime, subdir: str = None, filename: str = None, extension: str = None, no_timestamp: bool = False) -> Tuple[Path, Path]:
    '''
    Obsidian helper. Takes a datetime and optional subdirectory name, filename, and extension.
    If an extension is provided, it ensures the path is to a file with that extension.
    If no extension is provided, it treats the path as a directory.
    '''
    year = date_time.strftime(YEAR_FMT)
    month = date_time.strftime(MONTH_FMT)
    day = date_time.strftime(DAY_FMT)
    day_short = date_time.strftime(DAY_SHORT_FMT)
    timestamp = date_time.strftime("%H%M%S")

    relative_path = Path("journal") / year / month / day
    if not subdir and not filename and not extension:
        relative_path = relative_path / f"{day}.md"

    else:
        if subdir:
            relative_path = relative_path / f"{day_short} {subdir}"

        if filename:
            if extension:
                extension = extension if extension.startswith(".") else f".{extension}"
            else:
                extension = validate_extension(filename, [".md", ".m4a", ".wav", ".aiff", ".flac", ".mp3", ".mp4", ".pdf", ".js", ".json", ".yaml", ".py"]) or ".md"

            filename = sanitize_filename(filename)
            filename = f"{day_short} {filename}" if no_timestamp else f"{day_short} {timestamp} {filename}"
            filename = f"{filename}{extension}" if not filename.endswith(extension) else filename
            relative_path = relative_path / filename

        else:
            debug(f"This only happens, theoretically, when no filename nor subdirectory are provided, but an extension is. Which is kinda silly.")
            return None, None

    absolute_path = OBSIDIAN_VAULT_DIR / relative_path
    os.makedirs(absolute_path.parent, exist_ok=True)
    return absolute_path, relative_path


def validate_extension(filename, valid_extensions=None):
    if valid_extensions is None:
        return os.path.splitext(filename)
    else:
        extension = os.path.splitext(filename)[-1].lower()
        return extension if extension in valid_extensions else None

def prefix_lines(text: str, prefix: str = '> ') -> str:
    lines = text.split('\n')
    prefixed_lines = [f"{prefix}{line.lstrip()}" for line in lines]
    return '\n'.join(prefixed_lines)

def f(file):
    if hasattr(file, 'read') and callable(file.read):
        return file
    if isinstance(file, (bytes, bytearray)):
        return file

    if isinstance(file, Path):
        file_path = file
    elif isinstance(file, str):
        file_path = Path(file)
    else:
        raise TypeError("Invalid file type. Expected str, Path, or file-like object.")

    with open(file_path, 'rb') as thefile:
        return thefile


def is_ad_or_tracker(url: str, rules: AdblockRules) -> bool:
    parsed_url = urlparse(url)
    return rules.should_block(url, { 'domain': parsed_url.netloc })

            
def contains_blacklisted_word(text: str, blacklist: List[str]) -> bool:
        return any(word.lower() in text.lower() for word in blacklist)
    
    
def contains_profanity(content: str, threshold: float = 0.01, custom_words: Optional[List[str]] = None) -> bool:
    custom_words = custom_words or []
    
    # Combine the profanity library's word list with custom words
    profanity.load_censor_words(custom_words)
    
    word_list = content.split()
    content_profanity_count = sum(1 for word in word_list if profanity.contains_profanity(word))
    content_profanity_ratio = content_profanity_count / len(word_list) if word_list else 0
    
    debug(f"Profanity ratio for content: {content_profanity_ratio}")
    return content_profanity_ratio >= threshold


def load_filter_lists(blocklists_dir: Path):
        rules = []
        for file_path in blocklists_dir.glob('*.txt'):
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    rules.extend(file.read().splitlines())
                info(f"Loaded blocklist: {file_path.name}")
            except Exception as e:
                err(f"Error loading blocklist {file_path.name}: {str(e)}")
        return rules
    
    
def initialize_adblock_rules(blocklists_dir: Path):
    rules = load_filter_lists(blocklists_dir)
    info(f"Initialized AdblockRules with {len(rules)} rules")
    return AdblockRules(rules)


def get_extension(file):
    try:
        if isinstance(file, str):
            file_path = Path(file)
        elif isinstance(file, Path):
            file_path = file
        else:
            file_path = Path(file.filename)
        file_extension = file_path.suffix
        return file_extension

    except Exception as e:
        err(f"Unable to get extension of {file}")
        raise e



def sanitize_filename(text, extension: str = None, max_length: int = MAX_PATH_LENGTH):
    """Sanitize a string to be used as a safe filename while protecting the file extension."""
    debug(f"Filename before sanitization: {text}")

    text = re.sub(r'\s+', ' ', text)
    sanitized = re.sub(ALLOWED_FILENAME_CHARS, '', text)
    sanitized = sanitized.strip()
    base_name, extension = os.path.splitext(sanitized)

    max_base_length = max_length - len(extension)
    if len(base_name) > max_base_length:
        base_name = base_name[:max_base_length - 5].rstrip()
    final_filename = base_name + extension

    debug(f"Filename after sanitization: {final_filename}")
    return final_filename


def check_file_name(file_name, max_length=255):
    """Check if the file name needs sanitization based on the criteria of the second sanitize_filename function."""

    needs_sanitization = False

    if len(file_name) > max_length:
        debug(f"Filename exceeds maximum length of {max_length}: {file_name}")
        needs_sanitization = True
    if re.search(ALLOWED_FILENAME_CHARS, file_name):
        debug(f"Filename contains non-word characters (except space, dot, and hyphen): {file_name}")
        needs_sanitization = True
    if re.search(r'\s{2,}', file_name):
        debug(f"Filename contains multiple consecutive spaces: {file_name}")
        needs_sanitization = True
    if file_name != file_name.strip():
        debug(f"Filename has leading or trailing spaces: {file_name}")
        needs_sanitization = True

    return needs_sanitization


def bool_convert(value: str = Form(None)):
    return value.lower() in ["true", "1", "t", "y", "yes"]


def str_to_bool(value: str) -> bool:
    """
    Convert a string to a boolean.
    Interprets 'true', '1', 'yes', 'y' as True.
    Interprets 'false', '0', 'no', 'n', '', or any other string as False.
    """

def get_timestamp():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


async def extract_text(file_path: str) -> str:
    """Extract text from file."""
    if file_path.endswith('.pdf'):
        return await extract_text_from_pdf(file_path)

    elif file_path.endswith('.docx'):
        return await extract_text_from_docx(file_path)


def clean_text(text):
    text = text.replace('-', '')
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)
    return text.strip()


async def ocr_pdf(file_path: str) -> str:
    try:
        images = await asyncio.to_thread(convert_from_path, file_path)
        texts = await asyncio.gather(*(asyncio.to_thread(pytesseract.image_to_string, image) for image in images))
        return ' '.join(texts)
    except Exception as e:
        err(f"Error during OCR: {str(e)}")
        return ""


async def extract_text_from_pdf(file_path: str) -> str:
    if not await is_valid_pdf(file_path):
        err(f"Invalid PDF file: {file_path}")
        return ""

    text = ''
    num_pages = 0

    # First, attempt to extract text using PyPDF2
    try:
        reader = await asyncio.to_thread(PdfReader, file_path)
        for page in reader.pages:
            text_content = page.extract_text() + ' ' if page.extract_text() else ''
            text += text_content
        num_pages = len(reader.pages)

        # If text was extracted successfully and it's deemed sufficient, return it
        if text and not should_use_ocr(text, num_pages):
            return clean_text(text)
    except Exception as e:
        err(f"Error extracting text with PyPDF2: {str(e)}")

    # If PyPDF2 extraction fails or is insufficient, fall back to pdfminer.six
    try:
        text_pdfminer = await asyncio.to_thread(pdfminer_extract_text, file_path)
        if text_pdfminer and not should_use_ocr(text_pdfminer, num_pages):
            return clean_text(text_pdfminer)
    except Exception as e:
        err(f"Error extracting text with pdfminer.six: {e}")

    # If both methods fail or are deemed insufficient, use OCR as the last resort
    debug("Falling back to OCR for text extraction...")
    return await ocr_pdf(file_path)

async def is_valid_pdf(file_path: str) -> bool:
    """Check if the file at file_path is a valid PDF."""
    try:
        kind = filetype.guess(file_path)
        return kind.mime == 'application/pdf'
    except Exception as e:
        err(f"Error checking file type: {e}")
        return False

async def extract_text_from_pdf(file_path: str) -> str:
    if not await is_valid_pdf(file_path):
        err(f"Invalid PDF file: {file_path}")
        return ""

    text = ''
    try:
        reader = await asyncio.to_thread(PdfReader, file_path)
        for page in reader.pages:
            text_content = page.extract_text() + ' ' if page.extract_text() else ''
            text += text_content
        if text.strip():  # Successfully extracted text
            return clean_text(text)
    except Exception as e:
        err(f"Error extracting text with PyPDF2: {str(e)}")

    try:
        text_pdfminer = await asyncio.to_thread(pdfminer_extract_text, file_path)
        if text_pdfminer.strip():  # Successfully extracted text
            return clean_text(text_pdfminer)
    except Exception as e:
        err(f"Error extracting text with pdfminer.six: {str(e)}")

    # Fall back to OCR
    debug("Falling back to OCR for text extraction...")
    try:
        images = convert_from_path(file_path)
        ocr_texts = await asyncio.gather(*(asyncio.to_thread(pytesseract.image_to_string, img) for img in images))
        return ' '.join(ocr_texts).strip()
    except Exception as e:
        err(f"OCR failed: {str(e)}")
        return ""

async def extract_text_from_docx(file_path: str) -> str:
    def read_docx(file_path):
        doc = Document(file_path)
        full_text = [paragraph.text for paragraph in doc.paragraphs]
        return '\n'.join(full_text)

    return await asyncio.to_thread(read_docx, file_path)

# Correcting read_text_file to be asynchronous
async def read_text_file(file_path: str) -> str:
    # This opens and reads a file asynchronously by offloading to a separate thread
    return await asyncio.to_thread(_sync_read_text_file, file_path)

def _sync_read_text_file(file_path: str) -> str:
    # Actual synchronous file reading operation
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def should_use_ocr(text, num_pages) -> bool:
    if not text:
        return True  # No text was extracted, use OCR
    word_count = len(text.split())
    avg_words_per_page = word_count / num_pages
    return avg_words_per_page < 10


def convert_to_unix_time(iso_date_str):
    dt = parser.parse(iso_date_str)  # Automatically parses datetime with timezone
    return int(dt.timestamp())


def haversine(lat1, lon1, lat2, lon2):
    """ Calculate the great circle distance between two points on the earth specified in decimal degrees. """
    lat1, lon1, lat2, lon2 = map(math.radians, [lat1, lon1, lat2, lon2])
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    a = math.sin(dlat / 2) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2) ** 2
    c = 2 * math.asin(math.sqrt(a))
    r = 6371  # Radius of Earth in kilometers
    return c * r



def convert_degrees_to_cardinal(d):
    """
    Convert degrees to cardinal directions
    """
    dirs = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE", "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
    ix = round(d / (360. / len(dirs)))
    return dirs[ix % len(dirs)]



HOURLY_COLUMNS_MAPPING = {
    "12am": "00:00:00",
    "2am": "02:00:00",
    "4am": "04:00:00",
    "6am": "06:00:00",
    "8am": "08:00:00",
    "10am": "10:00:00",
    "12pm": "12:00:00",
    "2pm": "14:00:00",
    "4pm": "16:00:00",
    "6pm": "18:00:00",
    "8pm": "20:00:00",
    "10pm": "22:00:00",
}

def convert_to_12_hour_format(datetime_obj_or_str):
    if isinstance(datetime_obj_or_str, str):
        try:
            datetime_obj = datetime.strptime(datetime_obj_or_str, "%Y-%m-%d %H:%M:%S")
        except ValueError:
            try:
                datetime_obj = datetime.strptime(datetime_obj_or_str, "%H:%M:%S")
            except ValueError:
                return "Invalid datetime string format"
    elif isinstance(datetime_obj_or_str, time):
        datetime_obj_or_str = datetime_obj_or_str.strftime("%H:%M:%S")
    else:
        datetime_obj = datetime_obj_or_str

    if isinstance(datetime_obj_or_str, str):
        time24 = datetime_obj_or_str
    else:
        time24 = datetime_obj.strftime("%H:%M:%S")

    reverse_mapping = {v: k for k, v in HOURLY_COLUMNS_MAPPING.items()}
    return reverse_mapping.get(time24, "Invalid time")


def encode_image_to_base64(image_path):
    if os.path.exists(image_path):
        with Image.open(image_path) as image:
            output_buffer = BytesIO()
            image.save(output_buffer, format='JPEG')
            byte_data = output_buffer.getvalue()
            base64_str = base64.b64encode(byte_data).decode('utf-8')
        return base64_str
    else:
        debug(f"Error: File does not exist at {image_path}")

def resize_and_convert_image(image_path, max_size=2160, quality=80):
    with Image.open(image_path) as img:
        # Resize image
        ratio = max_size / max(img.size)
        new_size = tuple([int(x * ratio) for x in img.size])
        img = img.resize(new_size, Image.Resampling.LANCZOS)

        # Convert to jpg
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='JPEG', quality=quality)
        img_byte_arr = img_byte_arr.getvalue()

    return img_byte_arr


def load_geonames_data(path: str):
    columns = ['geonameid', 'name', 'asciiname', 'alternatenames',
               'latitude', 'longitude', 'feature_class', 'feature_code',
               'country_code', 'cc2', 'admin1_code', 'admin2_code', 'admin3_code',
               'admin4_code', 'population', 'elevation', 'dem', 'timezone', 'modification_date']

    data = pd.read_csv(
        path,
        sep='\t',
        header=None,
        names=columns,
        low_memory=False
    )

    return data

async def run_ssh_command(server, command):
    try:
        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        ssh.connect(server.ssh.host, username=server.ssh.user, password=server.ssh.password)
        stdin, stdout, stderr = ssh.exec_command(command)
        output = stdout.read().decode()
        error = stderr.read().decode()
        ssh.close()
        return output, error
    except Exception as e:
        err(f"SSH command failed for server {server.id}: {str(e)}")
        raise


async def html_to_markdown(url: str = None, source: str = None) -> Optional[str]:
    if source:
        html_content = source
    elif url:
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                html_content = await response.text()
    else:
        err(f"Unable to convert nothing to markdown.")
        return None
    
    # Use readability to extract the main content
    doc = ReadabilityDocument(html_content)
    cleaned_html = doc.summary()
    
    # Parse the cleaned HTML with BeautifulSoup for any additional processing
    soup = BeautifulSoup(cleaned_html, 'html.parser')
    
    # Remove any remaining unwanted elements
    for element in soup(['script', 'style']):
        element.decompose()
    
    # Convert to markdown
    markdown_content = md(str(soup), heading_style="ATX")
    
    return markdown_content
